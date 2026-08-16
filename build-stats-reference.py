# Builds the Word reference for what the custom-game logger can collect, split by whether the
# person running it is playing or spectating. Kept beside the logger it documents, because the
# two go out of date together.
#   python build-stats-reference.py
import os

import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

OUT = os.path.dirname(os.path.abspath(__file__))

INK = RGBColor(0x1A, 0x1F, 0x2B)
ACCENT = RGBColor(0x0F, 0x76, 0x6E)
WARM = RGBColor(0xB4, 0x53, 0x09)
MUTED = RGBColor(0x6B, 0x72, 0x80)
FONT = "Segoe UI"

TITLE = "LWG Custom Game Logger — What Each Route Can Collect"
SUBTITLE = "Every stat available when a player runs the logger, and when a spectator does"
DATELINE = "August 2026  ·  LWG Team Tool  ·  custom-game-logger.ps1"

INTRO = (
    "The logger takes one of two routes, and picks between them on its own. Someone playing the "
    "game has it in their own match history, which is a single rich end-of-game record. A "
    "spectator does not — their account never played it — so their route reads the live "
    "scoreboard out of the running game instead, and keeps the last snapshot before the game "
    "closes. The two never both fire on one machine."
)

LEGEND = [
    ("Captured", "Written to the database today."),
    ("Available", "Present in the source, not yet captured. An add, not a research project."),
    ("Not available", "Absent from that route entirely. Needs the other route, or another source."),
]

CAVEAT = (
    "The spectator figures come from Riot's Live Client Data API and the player figures from the "
    "League client's own match-history record. Neither is covered by Riot's public API guarantees, "
    "and field availability moves between patches. Nothing marked Available has been tested "
    "end-to-end from a real spectated game yet — the first spectated scrim is the test."
)

SPECTATOR = [
    ("Per player — direct", [
        ("Champion", "Captured", ""),
        ("Riot ID / summoner name", "Captured", "Used to match against the roster"),
        ("Team (blue / red)", "Captured", ""),
        ("Position", "Captured", "Top / Jungle / Mid / ADC / Support"),
        ("Level", "Captured", ""),
        ("Kills, deaths, assists", "Captured", ""),
        ("Creep score", "Captured", ""),
        ("Items", "Captured", "All seven slots"),
        ("Summoner spells", "Captured", "Both"),
        ("Ward score", "Available", "A live approximation, not the true vision score"),
        ("Runes", "Available", "Keystone, primary tree, secondary tree"),
        ("Skin", "Available", "Makes matching players to VOD footage easier"),
        ("Alive / dead and respawn timer", "Available", ""),
        ("Gold spent", "Available", "Summed from item prices. Excludes sold items and consumables"),
    ]),
    ("Game — direct", [
        ("Game mode and duration", "Captured", ""),
        ("Map name and terrain", "Available", "Which drake soul the map turned to"),
    ]),
    ("From the event feed", [
        ("First blood, tower, inhibitor, dragon, baron, herald", "Captured", "Derived per team"),
        ("Tower / inhibitor / dragon / baron / herald counts", "Captured", "Derived per team"),
        ("Timestamp on every event", "Available", "The single biggest gain over the player route"),
        ("Dragon type per kill", "Available", "Infernal, Cloud, Mountain, and so on"),
        ("Stolen flag on dragon, baron, herald", "Available", "The most reviewable moment in a scrim"),
        ("Every champion kill", "Available", "Killer, victim, assisters, time"),
        ("Multikills and aces", "Available", "With streak size and acing team"),
        ("Per-player turret and inhibitor kills", "Available", "From the killer name on the event"),
        ("Largest killing spree", "Available", "Derived from the kill sequence"),
        ("Inhibitor respawn timings", "Available", ""),
    ]),
    ("Derived by polling over time — spectator only", [
        ("CS at 10 / 15 / 20 minutes, and CS per minute", "Available", ""),
        ("CS difference against the opposing laner", "Available", "At any minute, not just in buckets"),
        ("Level curve and level leads", "Available", ""),
        ("Objective timestamps", "Available", "First drake at 8:30, baron at 24:10"),
        ("Kill-difference swing over time", "Available", "Shows where the game actually turned"),
        ("Item build order and completion timings", "Available", ""),
        ("Time spent dead, longest time alive", "Available", ""),
        ("Ward score curve", "Available", ""),
    ]),
    ("Not obtainable while spectating", [
        ("Damage to champions", "Not available", "Including the physical / magic / true split"),
        ("Damage taken, self-mitigated, to objectives, to turrets", "Not available", ""),
        ("Gold earned", "Not available", "Gold spent can be approximated from items"),
        ("Total healing and shielding", "Not available", ""),
        ("Time CCing others", "Not available", ""),
        ("Vision score, wards placed, wards killed, control wards", "Not available", "Only the live ward score exists"),
        ("Jungle CS split, own versus enemy jungle", "Not available", ""),
        ("Champion bans", "Not available", ""),
    ]),
]

PLAYER = [
    ("Per player — combat", [
        ("Kills, deaths, assists, win", "Captured", ""),
        ("Damage to champions", "Captured", "With the physical / magic / true split"),
        ("Damage taken and self-mitigated", "Captured", "Reads a tank's game the way damage dealt cannot"),
        ("Damage to objectives and turrets", "Captured", ""),
        ("Total healing", "Captured", ""),
        ("Time CCing others", "Captured", ""),
        ("Largest multikill and killing spree", "Captured", ""),
        ("Double / triple / quadra / penta kills", "Captured", ""),
        ("Total damage dealt", "Available", "Everything, not only to champions"),
        ("Largest critical strike", "Available", ""),
        ("Total units healed", "Available", "Separates a support's healing from a drain tank's"),
        ("Killing spree count, unreal kills", "Available", ""),
    ]),
    ("Per player — economy and farm", [
        ("Gold earned and gold spent", "Captured", ""),
        ("Total CS and lane CS", "Captured", ""),
        ("Own-jungle versus enemy-jungle CS", "Captured", "Farmed his own camps, or invaded yours"),
        ("Champion level", "Captured", ""),
    ]),
    ("Per player — vision", [
        ("Vision score", "Captured", ""),
        ("Wards placed and wards killed", "Captured", ""),
        ("Control wards bought", "Captured", ""),
        ("Sight wards bought", "Available", ""),
    ]),
    ("Per player — firsts and structures", [
        ("First blood kill, first tower kill", "Captured", ""),
        ("Turret kills, inhibitor kills", "Captured", ""),
        ("Longest time alive", "Captured", ""),
        ("First blood assist, first tower assist", "Available", ""),
        ("First inhibitor kill and assist", "Available", ""),
    ]),
    ("Per player — loadout", [
        ("Both summoner spells", "Captured", ""),
        ("Full seven-slot item build", "Captured", ""),
        ("Complete rune page", "Available", "Keystone, all six perks, both trees, stat shards"),
    ]),
    ("Per player — timeline deltas, in 10-minute buckets", [
        ("CS per minute", "Available", "0-10, 10-20, 20-30, 30+"),
        ("Gold per minute", "Available", ""),
        ("XP per minute", "Available", ""),
        ("Damage taken per minute", "Available", ""),
        ("CS difference against the lane opponent", "Available", "Laning phase, measured rather than argued"),
        ("XP difference against the lane opponent", "Available", ""),
        ("Damage-taken difference against the lane opponent", "Available", ""),
        ("Assigned lane and role", "Available", ""),
    ]),
    ("Team level", [
        ("First blood, tower, inhibitor, baron, dragon, herald", "Captured", ""),
        ("Tower / inhibitor / baron / dragon / herald counts", "Captured", ""),
        ("Champion bans", "Captured", "Customs are the one place bans are a real team decision"),
    ]),
    ("Game level", [
        ("Game id, duration, capture time", "Captured", ""),
        ("Patch version, queue id, map id, creation time", "Available", ""),
    ]),
    ("Not obtainable from a played game", [
        ("Event timestamps", "Not available", "The record has counts and who was first, never when"),
        ("Kill-by-kill sequence", "Not available", ""),
        ("Stolen-objective flags", "Not available", ""),
        ("Item purchase order", "Not available", "Only the final build"),
    ]),
]

CLOSING = [
    ("What each route is better at", [
        "The player route wins on magnitude: damage, gold, healing, crowd control, vision. None of "
        "it exists in the live feed.",
        "The spectator route wins on time: when objectives fell, how the lead moved, what order the "
        "build came together in. None of that is in the match-history record.",
        "They are complementary rather than redundant, so running both on one scrim is worth doing "
        "once the two records are de-duplicated.",
        "If only one person can run it, it should be someone playing — magnitude covers more review "
        "questions than timing does.",
    ]),
    ("What is worth adding next", [
        "Polling time series on the spectator route: objective timestamps, CS at 10/15/20, level and "
        "kill-difference curves. The largest single gain, and unavailable any other way.",
        "Stolen objectives and dragon types: a few lines of parsing against the event feed.",
        "Runes on both routes, and gold-spent-from-items on the spectator route.",
        "Timeline deltas on the player route: already in the payload, and the only measured read on "
        "laning phase either route can produce.",
    ]),
]


def style_document(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    for name, size, color in (("Heading 1", 18, INK), ("Heading 2", 13, ACCENT), ("Heading 3", 11.5, INK)):
        style = doc.styles[name]
        style.font.name = FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
    for section in doc.sections:
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)


def add_run(paragraph, text, size=10.5, color=INK, bold=False, italic=False):
    run = paragraph.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic
    return run


def add_stat_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = (Inches(3.1), Inches(1.15), Inches(2.75))

    header = table.rows[0].cells
    for cell, label in zip(header, ("Stat", "Status", "Notes")):
        cell.text = ""
        add_run(cell.paragraphs[0], label, size=9.5, color=ACCENT, bold=True)

    for stat, status, note in rows:
        cells = table.add_row().cells
        cells[0].text = ""
        add_run(cells[0].paragraphs[0], stat, size=9.5)
        cells[1].text = ""
        # Colour carries the status at a glance; the word is there for anyone printing in mono.
        colour = {"Captured": ACCENT, "Available": WARM}.get(status, MUTED)
        add_run(cells[1].paragraphs[0], status, size=9.5, color=colour, bold=(status == "Captured"))
        cells[2].text = ""
        add_run(cells[2].paragraphs[0], note, size=9, color=MUTED, italic=True)

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = width

    doc.add_paragraph()


def add_route(doc, heading, blurb, groups):
    doc.add_page_break()
    doc.add_heading(heading, level=1)
    para = doc.add_paragraph()
    add_run(para, blurb, size=10, color=MUTED, italic=True)
    for title, rows in groups:
        doc.add_heading(title, level=2)
        add_stat_table(doc, rows)


def build(path):
    doc = docx.Document()
    style_document(doc)

    title = doc.add_paragraph()
    add_run(title, TITLE, size=22, bold=True)
    subtitle = doc.add_paragraph()
    add_run(subtitle, SUBTITLE, size=12, color=MUTED)
    dateline = doc.add_paragraph()
    add_run(dateline, DATELINE, size=9, color=MUTED, italic=True)

    doc.add_paragraph()
    intro = doc.add_paragraph()
    add_run(intro, INTRO, size=10.5)

    doc.add_heading("How to read the status column", level=2)
    for label, meaning in LEGEND:
        para = doc.add_paragraph(style="List Bullet")
        colour = {"Captured": ACCENT, "Available": WARM}.get(label, MUTED)
        add_run(para, label + " — ", size=10, color=colour, bold=True)
        add_run(para, meaning, size=10)

    add_route(
        doc,
        "Spectator",
        "Read live from the running game on port 2999. Passive — no camera position, screen or "
        "player selection is required, because the feed carries whole-game state rather than "
        "what is on screen. The game window must stay open until the game ends.",
        SPECTATOR,
    )

    add_route(
        doc,
        "Player",
        "Read once from the client's own match-history record after the game. Any one of the ten "
        "players can run it — their record contains all ten, both teams.",
        PLAYER,
    )

    doc.add_page_break()
    doc.add_heading("Reading the two together", level=1)
    for title, points in CLOSING:
        doc.add_heading(title, level=2)
        for point in points:
            para = doc.add_paragraph(style="List Bullet")
            add_run(para, point, size=10)

    doc.add_paragraph()
    caveat = doc.add_paragraph()
    add_run(caveat, CAVEAT, size=9, color=WARM, italic=True)

    doc.save(path)
    return path


if __name__ == "__main__":
    out = build(os.path.join(OUT, "LWG-Logger-Stats-Reference.docx"))
    print("wrote", os.path.basename(out))
